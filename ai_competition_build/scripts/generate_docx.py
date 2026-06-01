from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
AI = next(p for p in ROOT.iterdir() if p.is_dir() and p.name.startswith("ai") and "竞赛资料" in p.name)
OUT = AI / "AI竞赛综合案例报告_大图连续浏览SDK与展示平台.docx"
LOGO = AI / "汉宁logo.png"


def find_dir(base: Path, *parts: str) -> Path:
    current = base
    for part in parts:
        current = next(p for p in current.iterdir() if p.is_dir() and part in p.name)
    return current


SDK_IMG = sorted(find_dir(AI, "大图连续浏览SDK参赛交付包", "软件界面图像").glob("*.png"))
PLATFORM_IMG = sorted(find_dir(AI, "展示平台参赛交付包", "软件界面图像").glob("*.png"))


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color="DCE6EF"):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:{}".format(edge)
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_text(cell, text, bold=False, color="18212F", size=9.5, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align:
        p.alignment = align
    r = p.add_run(str(text))
    r.bold = bold
    r.font.name = "Microsoft YaHei"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r.font.size = Pt(size)
    r.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True, color="FFFFFF", size=9)
        set_cell_shading(table.rows[0].cells[i], "0E2A3F")
        set_cell_border(table.rows[0].cells[i], "0E2A3F")
        if widths:
            table.rows[0].cells[i].width = Cm(widths[i])
    for r_idx, row in enumerate(rows):
        cells = table.add_row().cells
        for c_idx, value in enumerate(row):
            set_cell_text(cells[c_idx], value, size=8.7)
            set_cell_shading(cells[c_idx], "FFFFFF" if r_idx % 2 == 0 else "F6F9FB")
            set_cell_border(cells[c_idx])
            if widths:
                cells[c_idx].width = Cm(widths[c_idx])
    doc.add_paragraph()
    return table


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        run.font.color.rgb = RGBColor.from_string("0E2A3F" if level == 1 else "1D5D9B")
    return p


def add_para(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.18
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.name = "Microsoft YaHei"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        r.font.color.rgb = RGBColor.from_string("1D5D9B")
        text = text[len(bold_prefix):]
    r = p.add_run(text)
    r.font.name = "Microsoft YaHei"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r.font.size = Pt(10.5)
    r.font.color.rgb = RGBColor.from_string("18212F")
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(item)
        r.font.name = "Microsoft YaHei"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor.from_string("18212F")


def add_image(doc, img, caption, width=Cm(14.2)):
    if not img or not img.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run("【待补充图片】" + caption)
        r.font.name = "Microsoft YaHei"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        r.font.color.rgb = RGBColor.from_string("5F6B7A")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(img), width=width)
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    r.font.name = "Microsoft YaHei"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r.font.size = Pt(8.5)
    r.font.color.rgb = RGBColor.from_string("5F6B7A")


def configure_doc(doc):
    section = doc.sections[0]
    section.top_margin = Cm(1.6)
    section.bottom_margin = Cm(1.45)
    section.left_margin = Cm(1.75)
    section.right_margin = Cm(1.75)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("AI竞赛综合案例报告 - 大图连续浏览SDK与展示平台")
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string("5F6B7A")

    styles = doc.styles
    styles["Normal"].font.name = "Microsoft YaHei"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    styles["Normal"].font.size = Pt(10.5)


def cover(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(22)
    if LOGO.exists():
        p.add_run().add_picture(str(LOGO), width=Cm(3.2))

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(20)
    title.paragraph_format.space_after = Pt(10)
    r = title.add_run("AI 赋能研发交付综合案例报告")
    r.bold = True
    r.font.name = "Microsoft YaHei"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r.font.size = Pt(24)
    r.font.color.rgb = RGBColor.from_string("0E2A3F")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("大图连续浏览 SDK 与地铁隧道展示平台")
    r.font.name = "Microsoft YaHei"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r.font.size = Pt(15)
    r.font.color.rgb = RGBColor.from_string("1D5D9B")

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_before = Pt(28)
    r = meta.add_run("面向公司 AI 应用竞赛 | 2026")
    r.font.name = "Microsoft YaHei"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor.from_string("5F6B7A")

    add_table(
        doc,
        ["汇报单元", "项目定位", "核心成果"],
        [
            ["项目一", "大图连续浏览 SDK", "GB 级隧道影像切片、LOD 浏览、高清瓦片加载、病害标注"],
            ["项目二", "地铁隧道展示平台", "台账同步、统一库、WebAPI、可视化页面、Swagger 接口"],
        ],
        widths=[2.5, 4.2, 8.0],
    )
    doc.add_page_break()


def executive_summary(doc):
    add_heading(doc, "一、汇报摘要", 1)
    add_para(
        doc,
        "本材料把两个 AI 辅助项目按项目单元整合汇报：大图连续浏览 SDK 解决底层超大影像浏览、切片和标注能力；地铁隧道展示平台解决成果数据入库、接口服务和可视化展示能力。两者组合后，形成从检测影像预处理、交互浏览、病害标注到成果汇报展示的完整链路。",
    )
    add_para(
        doc,
        "材料结构严格对应竞赛要求：每个项目均包含 AI 工具与沟通过程、软件介绍/亮点/运行效果、量化指标、结论及体会。缺少真实 AI 对话截图的位置已用“待补充图片”占位，便于赛前替换为真实过程证据。",
    )
    add_table(
        doc,
        ["项目", "AI价值概括", "关键量化指标"],
        [
            ["大图连续浏览 SDK", "AI 参与架构拆分、Qt/C++ 实现、性能策略、调试与文档沉淀。", "212 个相关文件、13,024 行；传统 120h / AI 35h，效率提升约 71%（估算）。"],
            ["地铁隧道展示平台", "AI 参与编译诊断、数据库映射、接口、前端展示、报告和答辩材料。", "56 个源码文件、12,215 行、23 个 API、1,461 条病害记录；效率提升约 76%。"],
        ],
        widths=[3.3, 6.2, 5.2],
    )


def sdk_section(doc):
    add_heading(doc, "二、项目一：大图连续浏览 SDK", 1)
    add_heading(doc, "2.1 AI 工具与沟通过程", 2)
    add_para(doc, "主要工具：ChatGPT / Codex。补充工具包括 OpenCode 千问、MiniMax 等，用于局部 Bug 思路、替代方案和表达优化。")
    add_bullets(
        doc,
        [
            "场景澄清：向 AI 说明 GB 级隧道全景图浏览、缩放、标注和普通 PC 运行约束。",
            "架构拆解：把系统拆为 ImageSlicer、SDK 主体和 DemoApp 三部分，分别承担切片、核心浏览标注和演示验证。",
            "实现迭代：围绕 Qt/C++、LOD、动态瓦片加载、异步解码、JSON 标注保存等模块持续生成和修正代码。",
            "人工验收：通过编译、运行截图、真实界面效果和 README 说明进行复核，避免 AI 输出脱离工程环境。",
        ],
    )
    add_image(doc, None, "AI 工具沟通过程截图（待补充：可放 Codex/ChatGPT 对话、关键提示词和修改记录）", width=Cm(13.5))

    add_heading(doc, "2.2 软件介绍、亮点与运行效果", 2)
    add_para(
        doc,
        "大图连续浏览 SDK 基于 Qt 5.8+ 和 C++ 实现，面向超高分辨率隧道全景扫描图。系统采用类似地图软件的瓦片加载方式，只加载视野范围内的图像切片，并通过 LOD 机制在缩小时显示缩略图、放大时异步加载高清切片。",
    )
    add_bullets(
        doc,
        [
            "数据预处理：ImageSlicer 将原始大图切分为 thumbnail、info.json 和 x_y.jpg 瓦片。",
            "浏览体验：支持 WASD/方向键移动视野、鼠标拖拽、Ctrl+滚轮缩放。",
            "标注能力：支持折线和闭合多边形绘制，支持裂缝、渗水、剥落等病害类型记录。",
            "性能策略：后台线程读取与解码图片，主界面不因大图解码而卡顿。",
            "工程约束：README 明确要求 MSVC 64-bit 构建，避免 32-bit 内存限制导致大图处理崩溃。",
        ],
    )
    add_image(doc, SDK_IMG[0], "图 1：图像裁剪导入数据库软件运行效果", width=Cm(13.5))
    add_image(doc, SDK_IMG[1], "图 2：裁剪软件运行结果", width=Cm(13.2))
    add_image(doc, SDK_IMG[2], "图 3：缩略图运行界面", width=Cm(15.0))
    add_image(doc, SDK_IMG[3], "图 4：高清图像运行截图", width=Cm(15.0))

    add_heading(doc, "2.3 量化指标", 2)
    add_table(
        doc,
        ["指标", "数值", "口径说明"],
        [
            ["相关源码/工程文件", "212 个", "统计 SDK、DemoApp、ImageSlicer 相关源码和工程文件；排除 .git、.vs、参赛资料和第三方 OpenCV 头文件。"],
            ["源码/工程行数", "13,024 行", "同上口径；其中 SDK 主体 7,229 行、DemoApp 3,471 行、ImageSlicer 2,272 行。"],
            ["运行截图证据", "4 张", "覆盖切片导入、裁剪结果、缩略图浏览和高清图浏览。"],
            ["效率提升估算", "约 71%", "传统开发 120h / AI 协作 35h，节省约 85h；建议赛前按真实记录校准。"],
            ["可复用模块", "3 个", "ImageSlicer、SDK 核心、DemoApp。"],
        ],
        widths=[3.4, 3.0, 8.3],
    )
    add_heading(doc, "2.4 结论及体会", 2)
    add_para(
        doc,
        "结论：大图连续浏览 SDK 已形成可演示的切片预处理、缩略图浏览、高清瓦片加载和病害标注能力，可作为地铁隧道检测成果展示和后续平台集成的底层图像能力。",
        bold_prefix="结论：",
    )
    add_para(
        doc,
        "体会：AI 的价值不只是生成代码，而是帮助把性能策略、工程结构、调试经验和说明文档串成闭环。有效方式是持续提供真实约束，例如数据规模、构建环境、运行截图和错误信息，再由人工进行最终验收。",
        bold_prefix="体会：",
    )


def platform_section(doc):
    add_heading(doc, "三、项目二：地铁隧道展示平台", 1)
    add_heading(doc, "3.1 AI 工具与沟通过程", 2)
    add_para(
        doc,
        "展示平台项目已有参赛交付包，本报告提取其中案例报告、效果数据表、Codex 使用日志和软件截图作为基础证据。AI 协作从“编译失败、数据表不一致、缺少直观展示界面”的状态出发，逐步形成上传工具、统一库、接口和 Web 展示闭环。",
    )
    add_bullets(
        doc,
        [
            "Codex：阅读代码、定位影响面，修改实体、初始化 SQL、DTO、Controller、Service 和前端页面。",
            "ChatGPT / Gemini：用于界面构思、方案讨论、材料表达和风险治理补充。",
            "人工输入：真实报错、业务规则、概要设计、样例目录、台账和验收反馈。",
            "验证闭环：通过 build、样例导入、Swagger、页面截图和文档输出减少“看起来能跑”的风险。",
        ],
    )
    add_image(doc, None, "AI 工具沟通过程截图（待补充：可放 Codex 使用日志、ChatGPT/Gemini 对话截图）", width=Cm(13.5))

    add_heading(doc, "3.2 软件介绍、亮点与运行效果", 2)
    add_para(
        doc,
        "展示平台面向地铁隧道检测成果的数据整理、入库、查询和展示。平台将台账、二维/三维 SQLite、灰度图、病害高清图和点云目录转化为统一库、接口服务和可视化展示能力，降低出差汇报、客户演示和成果验收中的环境与操作成本。",
    )
    add_bullets(
        doc,
        [
            "WinForms 上传工具：读取台账，同步工程实例，按站点/区间上传成果目录。",
            "统一库：保存病害、环片、灰度图、高清图、点云文件等结构化索引。",
            "WebAPI：补齐工程概览、站点区间、里程范围、病害统计、图片和文件树接口。",
            "Web 展示：支持工程选择、二维图连续浏览、病害列表、高清图弹窗、Swagger 接口文档。",
        ],
    )
    add_image(doc, PLATFORM_IMG[1], "图 5：隧道平台上传工具", width=Cm(14.0))
    add_image(doc, PLATFORM_IMG[2], "图 6：前端登录界面", width=Cm(15.0))
    add_image(doc, PLATFORM_IMG[3], "图 7：隧道项目信息界面", width=Cm(15.0))
    add_image(doc, PLATFORM_IMG[4], "图 8：隧道区间浏览界面", width=Cm(15.0))
    add_image(doc, PLATFORM_IMG[5], "图 9：Swagger 控制器接口", width=Cm(15.0))

    add_heading(doc, "3.3 量化指标", 2)
    add_table(
        doc,
        ["指标", "数值", "口径说明"],
        [
            ["效率提升", "约 76%", "传统估算 162h / AI 协作 39h，节省约 123h。"],
            ["源码文件", "56 个", "既有效果数据表口径：不含 bin/obj/.git/.vs/server-storage。"],
            ["源码行数", "12,215 行", "既有效果数据表口径：不含日志和运行输出。"],
            ["API 接口", "23 个", "Controller 中公开 HTTP 接口。"],
            ["病害记录", "1,461 条", "DISEASE_CHK_DATA 记录数，进入统一查询链路。"],
            ["环片记录", "1,258 条", "RING_LOC 记录数。"],
            ["灰度图索引", "34 条", "IMAGE_DATA 中灰度图记录。"],
            ["样例数据", "1.8GB / 763 个文件", "server-storage 中检测成果文件大小合计约值。"],
        ],
        widths=[3.1, 3.2, 8.4],
    )
    add_table(
        doc,
        ["工作环节", "传统估算", "AI协作", "AI价值说明"],
        [
            ["需求澄清与方案拆解", "8h", "1.5h", "生成任务拆分、风险清单和迭代路线。"],
            ["数据库结构设计与表映射", "24h", "5h", "按 Word 成果表形成统一库和索引。"],
            ["上传工具与台账同步", "20h", "4.5h", "快速迭代 WinForms 交互和字段兼容。"],
            ["SQLite 解析与入库逻辑", "40h", "11h", "识别 2D/3D 表映射、编码问题和导入边界。"],
            ["前端展示与交互联动", "36h", "9h", "生成大屏样式、图像浏览和病害联动。"],
            ["接口文档与材料整理", "16h", "3h", "自动整理 API 说明、流程图、竞赛报告和答辩材料。"],
            ["测试修复与日志沉淀", "18h", "5h", "根据构建错误、页面验证和业务反馈循环修复。"],
        ],
        widths=[4.0, 2.0, 2.0, 6.7],
    )
    add_heading(doc, "3.4 结论及体会", 2)
    add_para(
        doc,
        "结论：展示平台已形成从成果目录上传、统一库入库、接口查询到 Web 页面演示的基础闭环，适合在成果汇报、客户演示和内部验收场景中使用。",
        bold_prefix="结论：",
    )
    add_para(
        doc,
        "体会：AI 在跨模块定位、重复代码生成、材料沉淀方面效率明显，但涉及 SQLite 字段语义、Word 表设计和工程数据安全时，必须坚持“AI 加速、人工验收、真实数据验证”的原则。",
        bold_prefix="体会：",
    )


def final_section(doc):
    add_heading(doc, "四、综合风险治理与推广建议", 1)
    add_table(
        doc,
        ["风险类型", "项目表现", "控制措施", "竞赛表达"],
        [
            ["AI 幻觉/误映射", "字段、表语义和业务规则可能被 AI 推断错误。", "以 README、概要设计、真实表结构、编译和样例导入验证为准。", "AI 加速，人工验收把关。"],
            ["数据安全", "检测数据、台账和工程信息均可能敏感。", "优先本地运行和脱敏截图，赛前确认工程名称与单位名称可公开。", "适合企业内网/本地部署。"],
            ["质量一致性", "多轮 AI 修改容易造成接口、DTO、页面口径不一致。", "保留构建、接口调用、页面截图和文档清单。", "形成可复跑验证清单。"],
            ["过程证据不足", "部分 AI 对话记录分散在多台电脑和多个工具。", "赛前补充关键对话截图、工时台账和演示录屏。", "把个人试错变成团队资产。"],
        ],
        widths=[2.7, 4.2, 4.8, 3.0],
    )
    add_para(
        doc,
        "推广建议：这套方法适合复制到检测、运维、质量、安全等数据密集部门，尤其适合“多源文件 + 结构化入库 + 可视化展示 + 汇报材料沉淀”的场景。推广前建议准备标准提示词模板、脱敏规范、工时记录表、验证清单和可复用演示素材库。",
    )
    add_heading(doc, "五、赛前补充清单", 1)
    add_bullets(
        doc,
        [
            "补充 2-3 张 AI 对话截图：需求拆解、关键 Bug 定位、材料整理各一张。",
            "补充真实工时记录或由负责人确认当前复盘估算口径。",
            "确认截图、台账、日志和文档中工程名称、单位名称是否需要脱敏。",
            "如主办方允许视频，建议录制：SDK 缩放浏览、切片导入、展示平台上传、区间浏览、病害高清图、Swagger 接口。",
        ],
    )


def main():
    doc = Document()
    configure_doc(doc)
    cover(doc)
    executive_summary(doc)
    sdk_section(doc)
    platform_section(doc)
    final_section(doc)
    doc.core_properties.title = "AI竞赛综合案例报告：大图连续浏览SDK与展示平台"
    doc.core_properties.subject = "AI应用竞赛"
    doc.core_properties.author = "Codex"
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
